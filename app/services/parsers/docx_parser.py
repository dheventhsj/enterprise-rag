"""DOCX document parser."""

from io import BytesIO

from docx import Document as DocxDocument

from app.exceptions import ValidationError
from app.services.parsers.base import ParsedDocument, ParsedPage


class DOCXParser:
    """Extract text from Word documents."""

    def parse(self, content: bytes) -> ParsedDocument:
        """Parse DOCX bytes into structured pages (single section)."""
        try:
            document = DocxDocument(BytesIO(content))
        except Exception as exc:
            raise ValidationError("Failed to read DOCX file") from exc

        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        text = "\n".join(paragraphs)

        if not text:
            raise ValidationError("DOCX contains no extractable text")

        return ParsedDocument(pages=[ParsedPage(page_number=1, text=text)])
